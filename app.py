import streamlit as st
import pandas as pd
import numpy as np
import math
from collections import defaultdict
import plotly.express as px
import plotly.graph_objects as go
import warnings
import io
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import tempfile
import os
from PIL import Image
import plotly.io as pio
warnings.filterwarnings('ignore')

# 设置页面配置
st.set_page_config(
    page_title="标准化高斯2SFCA分析工具",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS样式
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        color: #2e86ab;
        margin-top: 2rem;
        margin-bottom: 1rem;
        border-bottom: 2px solid #2e86ab;
        padding-bottom: 0.5rem;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #d1ecf1;
        border: 1px solid #bee5eb;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .formula-box {
        background-color: #f8f9fa;
        border: 2px solid #dee2e6;
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1rem 0;
        font-family: 'Courier New', monospace;
    }
    .download-section {
        background-color: #e8f4fd;
        border: 1px solid #b6d7e8;
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

class NormalizedGaussian2SFCA:
    """标准化高斯两步移动搜索法分析类"""
    
    def __init__(self, l0_distance, cost_type="距离"):
        self.l0_distance = l0_distance
        self.cost_type = cost_type
        self._init_gaussian_constants()
        
    def _init_gaussian_constants(self):
        """初始化用户公式的常数项"""
        if self.l0_distance <= 0:
            self._boundary_const = 0.0
            self._denominator = 1.0
            return
            
        self._boundary_const = math.exp(-0.5)
        self._denominator = 1 - self._boundary_const
        
    def gaussian_weight(self, distance):
        """标准化高斯权重函数"""
        if distance >= self.l0_distance:
            return 0.0
        if self._denominator <= 0:
            return 0.0
            
        ratio_squared = (distance / self.l0_distance) ** 2
        weight_unnormalized = math.exp(-0.5 * ratio_squared)
        numerator = weight_unnormalized - self._boundary_const
        weight = numerator / self._denominator
        
        return max(0.0, weight)
    
    def calculate_accessibility(self, df):
        """计算可达性得分"""
        df = df.copy()
        df['UserFormulaWeight'] = df['TravelCost'].apply(self.gaussian_weight)
        
        # 计算每个供给点的加权需求
        supply_demand = defaultdict(float)
        for _, row in df.iterrows():
            if row['UserFormulaWeight'] > 0:
                supply_demand[row['SupplyID']] += row['Demand'] * row['UserFormulaWeight']
        
        # 计算供给比率
        supply_ratios = {}
        supply_data = df[['SupplyID', 'Supply']].drop_duplicates()
        for _, row in supply_data.iterrows():
            supply_id = row['SupplyID']
            supply_value = row['Supply']
            weighted_demand = supply_demand.get(supply_id, 0)
            supply_ratios[supply_id] = supply_value / weighted_demand if weighted_demand > 0 else 0
        
        # 计算每个需求点的可达性
        accessibility_scores = defaultdict(float)
        for _, row in df.iterrows():
            if row['UserFormulaWeight'] > 0:
                origin_id = row['DemandID']
                dest_id = row['SupplyID']
                weight = row['UserFormulaWeight']
                if dest_id in supply_ratios:
                    accessibility_scores[origin_id] += supply_ratios[dest_id] * weight
        
        # 创建结果DataFrame
        demand_points = df[['DemandID', 'Demand']].drop_duplicates()
        results = []
        for _, row in demand_points.iterrows():
            demand_id = row['DemandID']
            results.append({
                'DemandID': demand_id,
                'Demand': row['Demand'],
                'AccessibilityScore': accessibility_scores.get(demand_id, 0)
            })
        
        return pd.DataFrame(results), df, supply_ratios

def create_sample_data():
    """创建示例数据"""
    sample_data = {
        'DemandID': [3, 59, 29, 80, 131, 39, 132, 197, 99],
        'Demand': [35023, 36080, 24316, 26139, 41445, 34871, 24155, 28886, 45856],
        'SupplyID': [215, 215, 215, 215, 215, 215, 215, 215, 215],
        'Supply': [437, 437, 437, 437, 437, 437, 437, 437, 437],
        'TravelCost': [0.05, 8.95, 9.27, 10, 10.98, 12.77, 13.23, 13.68, 14.27]
    }
    return pd.DataFrame(sample_data)

def plot_gaussian_decay(l0_distance, cost_type):
    """绘制高斯衰减曲线"""
    distances = np.linspace(0, l0_distance * 1.5, 100)
    analyzer = NormalizedGaussian2SFCA(l0_distance, cost_type)
    weights = [analyzer.gaussian_weight(d) for d in distances]
    
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=distances, y=weights, mode='lines', name='标准化高斯权重',
        line=dict(color='#FF4B4B', width=3)
    ))
    
    fig.add_vline(x=l0_distance, line_dash="dash", line_color="blue", 
                  annotation_text=f"截止距离 l_0 = {l0_distance}")
    
    fig.update_layout(
        title=f'标准化高斯衰减函数 (l_0 = {l0_distance})',
        xaxis_title=f'{cost_type}',
        yaxis_title='权重',
        showlegend=True,
        height=400,
        template="plotly_white"
    )
    
    return fig

def plot_accessibility_distribution(results_df):
    """绘制可达性分布"""
    fig = px.histogram(
        results_df, x='AccessibilityScore', title='可达性得分分布',
        nbins=20, color_discrete_sequence=['#1f77b4']
    )
    fig.update_layout(
        xaxis_title='可达性得分', yaxis_title='频数', height=400,
        template="plotly_white"
    )
    return fig

def plot_od_connections(df_with_weights, cost_type):
    """绘制OD连接权重分布"""
    fig = px.scatter(
        df_with_weights, x='TravelCost', y='UserFormulaWeight',
        size='Demand', color='UserFormulaWeight', title='OD连接权重分布',
        hover_data=['DemandID', 'SupplyID'], color_continuous_scale='viridis'
    )
    fig.update_layout(
        xaxis_title=f'{cost_type}成本', yaxis_title='用户公式权重', 
        height=400, template="plotly_white"
    )
    return fig

def plot_accessibility_boxplot(results_df):
    """绘制可达性得分箱线图"""
    fig = px.box(
        results_df, y='AccessibilityScore', title='可达性得分分布箱线图',
        color_discrete_sequence=['#2ca02c']
    )
    fig.update_layout(
        yaxis_title='可达性得分', height=400, template="plotly_white"
    )
    return fig

def plot_accessibility_vs_demand(results_df):
    """绘制可达性vs需求散点图"""
    # 移除对statsmodels的依赖，使用简单的线性趋势线
    fig = px.scatter(
        results_df, x='Demand', y='AccessibilityScore',
        title='需求量与可达性关系',
        trendline="ols",  # 使用普通的线性回归
        color='AccessibilityScore',
        color_continuous_scale='viridis'
    )
    fig.update_layout(
        xaxis_title='需求量', yaxis_title='可达性得分', height=400,
        template="plotly_white"
    )
    return fig

def plot_accessibility_heatmap(results_df, df_with_weights):
    """绘制可达性热力图（替代TOP10排名）"""
    # 创建需求点-可达性得分的分布热力图
    fig = px.density_heatmap(
        results_df, x='Demand', y='AccessibilityScore',
        title='需求量与可达性关系热力图',
        nbinsx=20, nbinsy=20,
        color_continuous_scale='viridis'
    )
    fig.update_layout(
        xaxis_title='需求量', yaxis_title='可达性得分', height=400,
        template="plotly_white"
    )
    return fig

def create_word_report(results_df, df_with_weights, supply_ratios, analyzer, cost_type, cost_unit, l0_distance, 
                      fig_decay, fig_dist, fig_od, fig_box, fig_scatter, fig_heatmap):
    """创建Word格式分析报告"""
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    # 标题
    title = doc.add_heading('标准化高斯2SFCA可达性分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('一、分析基本信息', level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid'
    info_table.cell(0, 0).text = '分析时间'
    info_table.cell(0, 1).text = pd.Timestamp.now().strftime('%Y年%m月%d日 %H:%M:%S')
    info_table.cell(1, 0).text = '成本类型'
    info_table.cell(1, 1).text = cost_type
    info_table.cell(2, 0).text = '截止距离 l₀'
    info_table.cell(2, 1).text = f'{l0_distance} {cost_unit}'
    info_table.cell(3, 0).text = '需求点数量'
    info_table.cell(3, 1).text = f'{len(results_df)} 个'
    info_table.cell(4, 0).text = '供给点数量'
    info_table.cell(4, 1).text = f'{len(supply_ratios)} 个'
    
    # 第二部分：原理与方法
    doc.add_heading('二、分析方法与原理', level=1)
    
    # 方法介绍
    doc.add_heading('2.1 标准化高斯两步移动搜索法', level=2)
    method_para = doc.add_paragraph()
    method_para.add_run('标准化高斯两步移动搜索法(2SFCA)是一种改进的空间可达性分析方法，')
    method_para.add_run('通过引入标准化高斯衰减函数，更准确地模拟空间相互作用的衰减模式。')
    
    # 核心公式
    doc.add_heading('2.2 核心权重函数', level=2)
    formula_para = doc.add_paragraph()
    formula_para.add_run('本分析采用以下标准化高斯权重函数：\n').bold = True
    
    formula_text = doc.add_paragraph()
    formula_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_run = formula_text.add_run('S(l_rn) = [e^(-1/2 × (l_rn/l_0)²) - e^(-1/2)] / [1 - e^(-1/2)]  当 l_rn < l_0\n')
    formula_run.bold = True
    formula_text.add_run('S(l_rn) = 0                                                   当 l_rn ≥ l_0')
    
    # 公式参数说明
    doc.add_heading('2.3 公式参数说明', level=2)
    param_table = doc.add_table(rows=5, cols=2)
    param_table.style = 'Light Grid'
    param_table.cell(0, 0).text = '参数符号'
    param_table.cell(0, 1).text = '参数含义'
    param_table.cell(1, 0).text = 'l_rn'
    param_table.cell(1, 1).text = f'从需求点r到供给点n的{cost_type}成本'
    param_table.cell(2, 0).text = 'l_0'
    param_table.cell(2, 1).text = f'截止{cost_type}参数，决定空间相互作用的最大范围'
    param_table.cell(3, 0).text = 'e^(-1/2)'
    param_table.cell(3, 1).text = '边界常数，约等于0.6065'
    param_table.cell(4, 0).text = '1 - e^(-1/2)'
    param_table.cell(4, 1).text = '标准化分母，约等于0.3935'
    
    # 计算步骤
    doc.add_heading('2.4 计算步骤', level=2)
    steps = [
        "第一步：计算每个供给点的服务范围内所有需求点的加权需求",
        "第二步：计算每个供给点的供给比率（供给量/加权需求）", 
        "第三步：计算每个需求点的可达性得分（所有可达供给点的供给比率加权和）"
    ]
    
    for i, step in enumerate(steps, 1):
        step_para = doc.add_paragraph()
        step_para.add_run(f'步骤{i}：').bold = True
        step_para.add_run(step)
    
    # 方法优势
    doc.add_heading('2.5 方法优势', level=2)
    advantages = [
        "• 单参数控制：只需设置截止距离l_0，操作简便",
        "• 天然标准化：权重函数自动归一化到[0,1]区间",
        "• 边界平滑：在截止距离处连续平滑衰减，避免突变",
        "• 空间衰减合理：更符合实际的空间相互作用模式"
    ]
    
    for advantage in advantages:
        doc.add_paragraph(advantage)
    
    # 第三部分：分析结果
    doc.add_heading('三、可达性分析结果', level=1)
    
    # 统计摘要
    doc.add_heading('3.1 统计摘要', level=2)
    stats = results_df['AccessibilityScore'].describe()
    stats_table = doc.add_table(rows=8, cols=2)
    stats_table.style = 'Light Grid'
    stats_table.cell(0, 0).text = '统计指标'
    stats_table.cell(0, 1).text = '数值'
    stats_table.cell(1, 0).text = '平均值'
    stats_table.cell(1, 1).text = f"{stats['mean']:.6f}"
    stats_table.cell(2, 0).text = '最大值'
    stats_table.cell(2, 1).text = f"{stats['max']:.6f}"
    stats_table.cell(3, 0).text = '最小值'
    stats_table.cell(3, 1).text = f"{stats['min']:.6f}"
    stats_table.cell(4, 0).text = '标准差'
    stats_table.cell(4, 1).text = f"{stats['std']:.6f}"
    stats_table.cell(5, 0).text = '25%分位数'
    stats_table.cell(5, 1).text = f"{stats['25%']:.6f}"
    stats_table.cell(6, 0).text = '50%分位数'
    stats_table.cell(6, 1).text = f"{stats['50%']:.6f}"
    stats_table.cell(7, 0).text = '75%分位数'
    stats_table.cell(7, 1).text = f"{stats['75%']:.6f}"
    
    # 可达性分布
    doc.add_heading('3.2 可达性分布可视化', level=2)
    doc.add_paragraph('以下图表展示了本次可达性分析的详细结果：')
    
    # 插入图表 - 高斯衰减函数
    doc.add_heading('高斯衰减函数', level=3)
    decay_img = fig_to_image(fig_decay)
    doc.add_picture(decay_img, width=Inches(6))
    doc.add_paragraph('图1: 标准化高斯衰减函数曲线，显示权重随距离增加而衰减的模式')
    
    # 插入图表 - 可达性分布直方图
    doc.add_heading('可达性得分分布', level=3)
    dist_img = fig_to_image(fig_dist)
    doc.add_picture(dist_img, width=Inches(6))
    doc.add_paragraph('图2: 可达性得分频率分布直方图')
    
    # 插入图表 - OD连接权重分布
    doc.add_heading('OD连接权重分布', level=3)
    od_img = fig_to_image(fig_od)
    doc.add_picture(od_img, width=Inches(6))
    doc.add_paragraph('图3: OD连接权重与出行成本关系散点图')
    
    # 插入图表 - 箱线图
    doc.add_heading('可达性得分分布箱线图', level=3)
    box_img = fig_to_image(fig_box)
    doc.add_picture(box_img, width=Inches(6))
    doc.add_paragraph('图4: 可达性得分的统计分布箱线图')
    
    # 插入图表 - 散点图
    doc.add_heading('需求量与可达性关系', level=3)
    scatter_img = fig_to_image(fig_scatter)
    doc.add_picture(scatter_img, width=Inches(6))
    doc.add_paragraph('图5: 需求量与可达性得分关系散点图')
    
    # 插入图表 - 热力图
    doc.add_heading('需求量与可达性关系热力图', level=3)
    heatmap_img = fig_to_image(fig_heatmap)
    doc.add_picture(heatmap_img, width=Inches(6))
    doc.add_paragraph('图6: 需求量与可达性得分关系热力图')
    
    # 前10名可达性得分
    doc.add_heading('3.3 可达性得分排名前10', level=2)
    top_10 = results_df.nlargest(10, 'AccessibilityScore')
    rank_table = doc.add_table(rows=11, cols=4)
    rank_table.style = 'Light Grid'
    rank_table.cell(0, 0).text = '排名'
    rank_table.cell(0, 1).text = '需求点ID'
    rank_table.cell(0, 2).text = '需求量'
    rank_table.cell(0, 3).text = '可达性得分'
    
    for i, (_, row) in enumerate(top_10.iterrows(), 1):
        rank_table.cell(i, 0).text = str(i)
        rank_table.cell(i, 1).text = str(row['DemandID'])
        rank_table.cell(i, 2).text = str(row['Demand'])
        rank_table.cell(i, 3).text = f"{row['AccessibilityScore']:.6f}"
    
    # 供给比率
    doc.add_heading('3.4 供给点服务比率', level=2)
    supply_table = doc.add_table(rows=len(supply_ratios)+1, cols=2)
    supply_table.style = 'Light Grid'
    supply_table.cell(0, 0).text = '供给点ID'
    supply_table.cell(0, 1).text = '供给比率'
    
    for i, (supply_id, ratio) in enumerate(supply_ratios.items(), 1):
        supply_table.cell(i, 0).text = str(supply_id)
        supply_table.cell(i, 1).text = f"{ratio:.6f}"
    
    # 分析结论
    doc.add_heading('3.5 分析结论与建议', level=2)
    zero_count = (results_df['AccessibilityScore'] == 0).sum()
    conclusion = f"""
本次空间可达性分析基于标准化高斯2SFCA方法，使用截止距离{l0_distance}{cost_unit}。
共分析了{len(results_df)}个需求点和{len(supply_ratios)}个供给点的空间可达性分布情况。

主要发现：
1. 平均可达性得分为 {stats['mean']:.6f}，反映了整体服务水平。
2. 可达性得分范围为 [{stats['min']:.6f}, {stats['max']:.6f}]，差异系数为 {stats['std']/stats['mean']:.2f}。
3. 共有 {zero_count} 个需求点({zero_count/len(results_df)*100:.1f}%)的可达性得分为0，表明这些区域服务覆盖不足。
4. 空间可达性分布{"较为均匀" if stats['std'] < stats['mean'] * 0.3 else "存在明显差异"}。

建议措施：
• 对低可达性区域({zero_count}个零得分点)进行重点分析，考虑增加服务设施。
• 优化高需求区域的服务资源配置，提高服务效率。
• 定期监测可达性变化，评估政策干预效果。
"""
    doc.add_paragraph(conclusion)
    
    # 保存到字节流
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def fig_to_image(fig):
    """将Plotly图形转换为图片字节流"""
    img_bytes = pio.to_image(fig, format='png', width=800, height=400, scale=2)
    return io.BytesIO(img_bytes)

def display_formula_explanation():
    """显示详细的公式解释"""
    st.markdown("""
    ## 📐 标准化高斯两步移动搜索法公式详解

    ### 核心权重函数
    """)
    
    # 使用LaTeX格式显示公式
    st.latex(r"""
    S(l_{rn}) = 
    \begin{cases}
    \frac{e^{-\frac{1}{2} \left( \frac{l_{rn}}{l_0} \right)^2} - e^{-\frac{1}{2}}}{1 - e^{-\frac{1}{2}}}, & \text{如果 } l_{rn} < l_0 \\
    0, & \text{如果 } l_{rn} \geq l_0
    \end{cases}
    """)
    
    st.markdown("""
    ### 🧮 公式参数说明

    | 参数 | 描述 | 示例值 |
    |------|------|--------|
    | **$l_{rn}$** | 从需求点 $r$ 到供给点 $n$ 的出行成本 | 5.3分钟 / 800米 |
    | **$l_0$** | 截止距离参数，决定空间相互作用的范围 | 15分钟 / 1000米 |
    | **$e^{-\\frac{1}{2}}$** | 边界常数，约等于 0.6065 | 固定值 |
    | **$1 - e^{-\\frac{1}{2}}$** | 标准化分母，约等于 0.3935 | 固定值 |

    ### 📊 计算步骤详解

    #### 第一步：计算每个供给点的加权需求
    """)
    
    st.latex(r"P_j = \sum_{k \in \{d_{kj} \leq l_0\}} P_k \cdot S(d_{kj})")
    
    st.markdown("""
    - $P_k$: 需求点 $k$ 的人口/需求量
    - $S(d_{kj})$: 标准化高斯权重函数
    - 只考虑距离在 $l_0$ 范围内的需求点

    #### 第二步：计算每个供给点的供给比率
    """)
    
    st.latex(r"R_j = \frac{S_j}{P_j}")
    
    st.markdown("""
    - $S_j$: 供给点 $j$ 的服务能力
    - $P_j$: 第一步计算的加权需求
    - 比率表示单位需求分配到的服务资源

    #### 第三步：计算每个需求点的可达性得分
    """)
    
    st.latex(r"A_i = \sum_{j \in \{d_{ij} \leq l_0\}} R_j \cdot S(d_{ij})")
    
    st.markdown("""
    - $R_j$: 第二步计算的供给比率
    - $S(d_{ij})$: 标准化高斯权重函数
    - 最终得分反映该需求点的综合可达性水平

    ### 🎯 方法特点

    - **单参数控制**: 只需设置截止距离 $l_0$
    - **天然标准化**: 权重范围自动归一化到 [0, 1]
    - **边界平滑**: 在 $l_0$ 处连续平滑衰减到0
    - **空间衰减**: 更符合实际的空间相互作用模式
    """)

def main():
    """主应用函数"""
    
    # 应用标题和介绍
    st.markdown('<h1 class="main-header">🏥 标准化高斯2SFCA可达性分析工具</h1>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
        <div class="info-box">
        <b>📊 工具介绍：</b> 基于标准化高斯两步移动搜索法的空间可达性分析工具，
        支持CSV、Excel等多种数据格式，提供专业的可达性分析和报告生成。
        </div>
        """, unsafe_allow_html=True)
    
    # 公式解释页面 - 放在主内容区
    with st.expander("📐 点击查看详细公式和算法说明", expanded=False):
        display_formula_explanation()
    
    # 侧边栏 - 参数配置
    with st.sidebar:
        st.header("⚙️ 分析配置")
        
        # 成本类型选择
        cost_type = st.selectbox(
            "出行成本类型",
            ["距离", "时间"],
            help="选择TravelCost列的单位类型"
        )
        
        cost_unit = st.text_input(
            "成本单位",
            value="米" if cost_type == "距离" else "分钟",
            help="例如：米、公里、分钟、小时等"
        )
        
        st.markdown("---")
        st.subheader("📏 截止距离参数")
        
        st.markdown("""
        **l₀ 参数说明：**
        - 空间相互作用的最大范围
        - 超过此值的权重为0
        - 影响衰减曲线形状
        """)
        
        # 根据成本类型设置不同的默认值和范围
        if cost_type == "距离":
            default_l0, min_val, max_val = 15.0, 0.1, 10000.0
            presets = {"步行尺度 (800米)": 800, "自行车尺度 (3000米)": 3000, "驾车尺度 (10000米)": 10000}
        else:
            default_l0, min_val, max_val = 30.0, 0.1, 600.0
            presets = {"步行尺度 (15分钟)": 15, "自行车尺度 (30分钟)": 30, "驾车尺度 (60分钟)": 60}
        
        # 预设按钮
        selected_preset = st.selectbox("快速设置", ["自定义"] + list(presets.keys()))
        if selected_preset != "自定义":
            l0_distance = presets[selected_preset]
            st.info(f"已选择: {selected_preset}")
        else:
            l0_distance = st.slider(
                f"截止距离 l₀ ({cost_unit})",
                min_value=min_val, max_value=max_val, value=default_l0, step=1.0
            )
        
        st.markdown("---")
        st.subheader("📁 数据输入")
        
        uploaded_file = st.file_uploader(
            "上传数据文件",
            type=['csv', 'xlsx', 'xls'],
            help="支持CSV、Excel格式，应包含: DemandID, Demand, SupplyID, Supply, TravelCost"
        )
        
        if uploaded_file is not None:
            try:
                # 根据文件类型读取数据
                if uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(uploaded_file)
                else:  # Excel文件
                    df = pd.read_excel(uploaded_file)
                
                required_columns = ['DemandID', 'Demand', 'SupplyID', 'Supply', 'TravelCost']
                missing_columns = [col for col in required_columns if col not in df.columns]
                
                if missing_columns:
                    st.error(f"❌ 缺少列: {missing_columns}")
                    st.info("使用示例数据")
                    df = create_sample_data()
                else:
                    st.success("✅ 数据上传成功!")
            except Exception as e:
                st.error(f"❌ 文件读取错误: {str(e)}")
                st.info("使用示例数据")
                df = create_sample_data()
        else:
            st.info("使用示例数据")
            df = create_sample_data()
    
    # 主内容区 - 数据展示
    st.markdown('<div class="section-header">📋 数据概览</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.dataframe(df, use_container_width=True)
    
    with col2:
        st.markdown("""
        <div class="metric-card">
        <h3>📊 数据统计</h3>
        """, unsafe_allow_html=True)
        st.write(f"**需求点数量:** {df['DemandID'].nunique()}")
        st.write(f"**供给点数量:** {df['SupplyID'].nunique()}")
        st.write(f"**OD连接数量:** {len(df)}")
        st.write(f"**平均{cost_type}成本:** {df['TravelCost'].mean():.2f} {cost_unit}")
        st.write(f"**最大{cost_type}成本:** {df['TravelCost'].max():.2f} {cost_unit}")
        st.write(f"**最小{cost_type}成本:** {df['TravelCost'].min():.2f} {cost_unit}")
        st.markdown("</div>", unsafe_allow_html=True)
        
        st.markdown("""
        <div class="metric-card">
        <h3>⚙️ 分析参数</h3>
        """, unsafe_allow_html=True)
        st.write(f"**成本类型:** {cost_type}")
        st.write(f"**截止距离 l₀:** {l0_distance} {cost_unit}")
        st.markdown("</div>", unsafe_allow_html=True)
    
    # 分析按钮
    st.markdown("---")
    if st.button("🚀 开始可达性分析", type="primary", use_container_width=True):
        with st.spinner("正在进行可达性分析..."):
            try:
                # 执行分析
                analyzer = NormalizedGaussian2SFCA(l0_distance, cost_type)
                results_df, df_with_weights, supply_ratios = analyzer.calculate_accessibility(df)
                
                # 生成图表
                fig_decay = plot_gaussian_decay(l0_distance, cost_type)
                fig_dist = plot_accessibility_distribution(results_df)
                fig_od = plot_od_connections(df_with_weights, cost_type)
                fig_box = plot_accessibility_boxplot(results_df)
                fig_scatter = plot_accessibility_vs_demand(results_df)
                fig_heatmap = plot_accessibility_heatmap(results_df, df_with_weights)
                
                # 将结果存储在session state中，防止重新运行后消失
                st.session_state.results_df = results_df
                st.session_state.df_with_weights = df_with_weights
                st.session_state.supply_ratios = supply_ratios
                st.session_state.analyzer = analyzer
                st.session_state.fig_decay = fig_decay
                st.session_state.fig_dist = fig_dist
                st.session_state.fig_od = fig_od
                st.session_state.fig_box = fig_box
                st.session_state.fig_scatter = fig_scatter
                st.session_state.fig_heatmap = fig_heatmap
                st.session_state.analysis_complete = True
                st.session_state.cost_type = cost_type
                st.session_state.cost_unit = cost_unit
                st.session_state.l0_distance = l0_distance
                
                # 显示成功信息
                st.markdown("""
                <div class="success-box">
                <h3>✅ 分析完成！</h3>
                可达性分析已成功完成，以下是详细结果。
                </div>
                """, unsafe_allow_html=True)
                
            except Exception as e:
                st.error(f"❌ 分析过程中出现错误: {str(e)}")
                st.info("请检查数据格式和参数设置，或联系技术支持")
    
    # 显示分析结果（如果分析已完成）
    if st.session_state.get('analysis_complete', False):
        results_df = st.session_state.results_df
        df_with_weights = st.session_state.df_with_weights
        supply_ratios = st.session_state.supply_ratios
        analyzer = st.session_state.analyzer
        fig_decay = st.session_state.fig_decay
        fig_dist = st.session_state.fig_dist
        fig_od = st.session_state.fig_od
        fig_box = st.session_state.fig_box
        fig_scatter = st.session_state.fig_scatter
        fig_heatmap = st.session_state.fig_heatmap
        cost_type = st.session_state.cost_type
        cost_unit = st.session_state.cost_unit
        l0_distance = st.session_state.l0_distance
        
        # 结果显示
        st.markdown('<div class="section-header">📈 分析结果</div>', unsafe_allow_html=True)
        
        # 统计摘要
        col1, col2, col3, col4 = st.columns(4)
        accessibility_scores = results_df['AccessibilityScore']
        
        with col1:
            st.metric("平均可达性", f"{accessibility_scores.mean():.6f}")
        with col2:
            st.metric("最大可达性", f"{accessibility_scores.max():.6f}")
        with col3:
            st.metric("最小可达性", f"{accessibility_scores.min():.6f}")
        with col4:
            zero_count = (accessibility_scores == 0).sum()
            st.metric("零可达性点", f"{zero_count}/{len(results_df)}")
        
        # 结果表格
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("🏆 可达性排名")
            display_df = results_df.sort_values('AccessibilityScore', ascending=False)
            display_df['排名'] = range(1, len(display_df) + 1)
            st.dataframe(display_df[['排名', 'DemandID', 'Demand', 'AccessibilityScore']], 
                       use_container_width=True)
        
        with col2:
            st.subheader("⚖️ 供给比率")
            supply_df = pd.DataFrame([
                {'供给点ID': k, '供给比率': v} 
                for k, v in supply_ratios.items()
            ])
            st.dataframe(supply_df, use_container_width=True)
        
        # 可视化分析
        st.markdown('<div class="section-header">📊 可视化分析</div>', unsafe_allow_html=True)
        
        # 第一行图表
        col1, col2 = st.columns(2)
        
        with col1:
            st.plotly_chart(fig_decay, use_container_width=True)
            
            # 图表下载按钮
            col1a, col1b = st.columns(2)
            with col1a:
                png_decay = pio.to_image(fig_decay, format='png', scale=2)
                st.download_button(
                    label="📥 下载PNG",
                    data=png_decay,
                    file_name="高斯衰减函数.png",
                    mime="image/png",
                    use_container_width=True
                )
            with col1b:
                pdf_decay = pio.to_image(fig_decay, format='pdf')
                st.download_button(
                    label="📥 下载PDF",
                    data=pdf_decay,
                    file_name="高斯衰减函数.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        
        with col2:
            st.plotly_chart(fig_dist, use_container_width=True)
            
            # 图表下载按钮
            col2a, col2b = st.columns(2)
            with col2a:
                png_dist = pio.to_image(fig_dist, format='png', scale=2)
                st.download_button(
                    label="📥 下载PNG",
                    data=png_dist,
                    file_name="可达性分布直方图.png",
                    mime="image/png",
                    use_container_width=True
                )
            with col2b:
                pdf_dist = pio.to_image(fig_dist, format='pdf')
                st.download_button(
                    label="📥 下载PDF",
                    data=pdf_dist,
                    file_name="可达性分布直方图.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        
        # 第二行图表
        col1, col2 = st.columns(2)
        
        with col1:
            st.plotly_chart(fig_od, use_container_width=True)
            
            # 图表下载按钮
            col1a, col1b = st.columns(2)
            with col1a:
                png_od = pio.to_image(fig_od, format='png', scale=2)
                st.download_button(
                    label="📥 下载PNG",
                    data=png_od,
                    file_name="OD连接权重分布.png",
                    mime="image/png",
                    use_container_width=True
                )
            with col1b:
                pdf_od = pio.to_image(fig_od, format='pdf')
                st.download_button(
                    label="📥 下载PDF",
                    data=pdf_od,
                    file_name="OD连接权重分布.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        
        with col2:
            st.plotly_chart(fig_box, use_container_width=True)
            
            # 图表下载按钮
            col2a, col2b = st.columns(2)
            with col2a:
                png_box = pio.to_image(fig_box, format='png', scale=2)
                st.download_button(
                    label="📥 下载PNG",
                    data=png_box,
                    file_name="可达性箱线图.png",
                    mime="image/png",
                    use_container_width=True
                )
            with col2b:
                pdf_box = pio.to_image(fig_box, format='pdf')
                st.download_button(
                    label="📥 下载PDF",
                    data=pdf_box,
                    file_name="可达性箱线图.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        
        # 第三行图表
        col1, col2 = st.columns(2)
        
        with col1:
            st.plotly_chart(fig_scatter, use_container_width=True)
            
            # 图表下载按钮
            col1a, col1b = st.columns(2)
            with col1a:
                png_scatter = pio.to_image(fig_scatter, format='png', scale=2)
                st.download_button(
                    label="📥 下载PNG",
                    data=png_scatter,
                    file_name="需求量与可达性关系.png",
                    mime="image/png",
                    use_container_width=True
                )
            with col1b:
                pdf_scatter = pio.to_image(fig_scatter, format='pdf')
                st.download_button(
                    label="📥 下载PDF",
                    data=pdf_scatter,
                    file_name="需求量与可达性关系.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        
        with col2:
            st.plotly_chart(fig_heatmap, use_container_width=True)
            
            # 图表下载按钮
            col2a, col2b = st.columns(2)
            with col2a:
                png_heatmap = pio.to_image(fig_heatmap, format='png', scale=2)
                st.download_button(
                    label="📥 下载PNG",
                    data=png_heatmap,
                    file_name="需求量与可达性热力图.png",
                    mime="image/png",
                    use_container_width=True
                )
            with col2b:
                pdf_heatmap = pio.to_image(fig_heatmap, format='pdf')
                st.download_button(
                    label="📥 下载PDF",
                    data=pdf_heatmap,
                    file_name="需求量与可达性热力图.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        
        # 技术细节
        with st.expander("🔬 技术细节", expanded=False):
            st.subheader("公式常数")
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**边界常数 e^(-1/2):** {analyzer._boundary_const:.6f}")
            with col2:
                st.write(f"**标准化分母:** {analyzer._denominator:.6f}")
            
            st.subheader("权重计算示例")
            test_data = []
            test_distances = [0, l0_distance*0.25, l0_distance*0.5, l0_distance*0.75, l0_distance]
            for dist in test_distances:
                weight = analyzer.gaussian_weight(dist)
                ratio = dist / l0_distance if l0_distance > 0 else 0
                test_data.append({
                    f'{cost_type}({cost_unit})': f"{dist:.2f}",
                    'l_rn/l_0': f"{ratio:.2f}",
                    '权重': f"{weight:.4f}"
                })
            st.table(pd.DataFrame(test_data))
        
        # 下载结果
        st.markdown("""
        <div class="download-section">
        <h3>💾 下载分析结果</h3>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # 可达性结果下载
            csv_results = results_df.to_csv(index=False)
            st.download_button(
                label="📥 下载可达性结果 (CSV)",
                data=csv_results,
                file_name=f"可达性分析结果_l0_{l0_distance}.csv",
                mime="text/csv",
                use_container_width=True,
                key="download_accessibility"
            )
        
        with col2:
            # 详细OD权重下载
            csv_od = df_with_weights.to_csv(index=False)
            st.download_button(
                label="📥 下载详细OD数据 (CSV)",
                data=csv_od,
                file_name=f"OD连接数据_l0_{l0_distance}.csv",
                mime="text/csv",
                use_container_width=True,
                key="download_od"
            )
        
        with col3:
            # Word报告下载
            try:
                doc_io = create_word_report(results_df, df_with_weights, supply_ratios, 
                                          analyzer, cost_type, cost_unit, l0_distance,
                                          fig_decay, fig_dist, fig_od, fig_box, fig_scatter, fig_heatmap)
                st.download_button(
                    label="📄 下载完整分析报告 (Word)",
                    data=doc_io.getvalue(),
                    file_name=f"可达性分析报告_l0_{l0_distance}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="download_report"
                )
            except Exception as e:
                st.error(f"生成Word报告时出错: {str(e)}")
                st.info("请确保已安装 python-docx 库")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # 使用说明
    with st.expander("📖 使用指南", expanded=False):
        st.markdown("""
        ### 🎯 使用步骤
        
        1. **准备数据**：确保数据文件包含以下列：
           - `DemandID` - 需求点唯一标识
           - `Demand` - 需求量（如人口数量）
           - `SupplyID` - 供给点唯一标识  
           - `Supply` - 供给量（如服务设施容量）
           - `TravelCost` - 出行成本（距离或时间）
        
        2. **设置参数**：
           - 选择出行成本类型（距离或时间）
           - 设置合适的截止距离 l₀
           - 上传数据文件（支持CSV、Excel格式）或使用示例数据
        
        3. **运行分析**：点击"开始可达性分析"按钮
        
        4. **查看结果**：分析结果包括：
           - 可达性得分表格和排名
           - 统计摘要
           - 多种可视化图表（支持PNG/PDF下载）
           - 可下载的结果文件和完整分析报告
        
        ### 📏 参数设定建议
        
        **截止距离 l₀ 设定：**
        - **距离类型**：步行800-1500米，驾车5000-15000米
        - **时间类型**：步行15-30分钟，驾车30-60分钟
        
        ### 📊 结果解释
        
        - **可达性得分**：数值越高表示可达性越好
        - **供给比率**：供给量与加权需求的比值
        - **权重衰减**：反映空间相互作用的衰减模式
        
        ### ❓ 常见问题
        
        **Q: 为什么有些点的可达性得分为0？**
        A: 这可能是因为该需求点与所有供给点的距离都超过了截止距离 l₀
        
        **Q: 如何选择合适的 l₀ 值？**
        A: 根据实际出行行为和研究目的选择，可参考预设值或进行敏感性分析
        """)

# 初始化session state
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False

if __name__ == "__main__":
    main()
